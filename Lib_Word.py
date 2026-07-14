

import io
import os
import streamlit as st
from docx import Document
from docx.oxml import OxmlElement
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn
from docx.shared import Inches, Pt

# LIBRARIES
from Lib_Image import disegna_punti_critici



def remove_internal_borders(cell, top=False, bottom=False):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    if top:
        top_el = OxmlElement('w:top')
        top_el.set(qn('w:val'), 'nil')
        tcBorders.append(top_el)
    if bottom:
        bottom_el = OxmlElement('w:bottom')
        bottom_el.set(qn('w:val'), 'nil')
        tcBorders.append(bottom_el)



@st.cache_data
def get_template_bytes():
    with open("Template.docx", "rb") as f:
        return f.read()
    

def genera_report_finale(storico, uploaded_files=None):
    
    # Creiamo un nuovo documento dai bytes cachati
    template_data = get_template_bytes()
    doc = Document(io.BytesIO(template_data))
    
    font_name = st.session_state.settings.get("font", "Arial")
    font_size = st.session_state.settings.get("size", 9)

    # --- FUNZIONE DI SOSTITUZIONE ROBUSTA ---
    def replace_text_in_doc(doc, key, value, font_name, font_size):
        placeholder = f"{{{{{key}}}}}"
        # Elaborazione paragrafi
        for para in doc.paragraphs:
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, value)
                for run in para.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
        # Elaborazione tabelle
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if placeholder in para.text:
                            para.text = para.text.replace(placeholder, value)
                            for run in para.runs:
                                run.font.name = font_name
                                run.font.size = Pt(font_size)

    # 1. Sostituzione Anagrafica (escludendo il segnaposto allegati se vuoi gestirlo a parte)
    if "anagrafica" in st.session_state:
        for key, value in st.session_state.anagrafica.items():
            if key != "allegati":
                replace_text_in_doc(doc, key, value, font_name, font_size)

    # 2. GESTIONE ALLEGATI (Logica estesa a Paragrafi e Tabelle)
    def gestisci_allegati_in_container(container):
        for p in container.paragraphs:
            if "{{allegati}}" in p.text:
                p.text = "" # Pulisce il tag
                
                # Recuperiamo gli allegati ottimizzati dallo stato
                allegati = st.session_state.get("allegati_ottimizzati", [])
                
                if allegati:
                    for f_data in allegati:
                        # Aggiunta nome file
                        run = p.add_run(f"- {f_data['name']}\n")
                        run.bold = True
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
                        
                        # Inserimento immagine compressa
                        if f_data['type'].startswith("image"):
                            try:
                                # Creiamo uno stream dai bytes già ottimizzati
                                img_stream = io.BytesIO(f_data['bytes'])
                                p.add_run().add_picture(img_stream, width=Inches(4.0))
                                p.add_run("\n")
                            except Exception as e:
                                st.error(f"Errore caricamento allegato {f_data['name']}: {e}")
                return True
        return False

    # Cerca nel corpo principale
    if not gestisci_allegati_in_container(doc):
        # Cerca in tutte le tabelle (come per gli altri campi)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if gestisci_allegati_in_container(cell):
                        break


    # 1. Trova il paragrafo segnaposto
    target_paragraph = None
    for p in doc.paragraphs:
        if "###TABELLA_ANALISI###" in p.text:
            target_paragraph = p
            break
            
    if target_paragraph:
        p_element = target_paragraph._element
        
        # 2. Crea la tabella
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        
        # Intestazione
        hdr = table.add_row().cells
        hdr[0].text = "ANALISI TECNICHE"

        for cell in hdr:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), 'D9D9D9') # Inserisci qui il codice esadecimale del colore
            tcPr.append(shd)

    
        # 3. Ciclo sullo storico
        for idx, data in enumerate(storico):

            img_path = data.get("img_path")
            if img_path and os.path.exists(img_path):
                with open(img_path, "rb") as f:
                    image_bytes = f.read()
            else:
                image_bytes = None

            # Recupera i punti aggiornati (inclusi quelli manipolati dall'utente)
            punti_totali = [p for img_data in data["report"].get("analisi_per_immagine", []) for p in img_data['punti_critici']]
            
            # Riga titolo con riassunto specifico
            riassunto_specifico = data["report"].get("riassunto_generale", "Analisi Tecnica")
            
            # Riga titolo
            row_t = table.add_row()
            for cell in row_t.cells:
                # 1. Imposta l'allineamento verticale
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
                # 2. Rimuovi margini interni del paragrafo che potrebbero dare l'illusione di decentramento
                for para in cell.paragraphs:
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
            
            # Formattazione FOTOGRAFIA X
            paragrafo_foto = row_t.cells[0].add_paragraph(f"FOTOGRAFIA {idx + 1}")
            run_foto = paragrafo_foto.runs[0]
            run_foto.font.name = font_name
            run_foto.font.size = Pt(font_size)
            run_foto.bold = True
            
            # Formattazione RIASSUNTO
            paragrafo_t = row_t.cells[1].add_paragraph(riassunto_specifico)
            run_t = paragrafo_t.runs[0]
            run_t.font.name = font_name
            run_t.font.size = Pt(font_size)
            run_t.bold = True
            
            
            # Riga contenuto
            row_c = table.add_row()
            
            # --- DISEGNO IMMAGINE CON TAG ---
            img_path = data.get("img_path")
            if img_path and os.path.exists(img_path):
                with open(img_path, "rb") as f:
                    image_bytes = f.read()
                
                img_taggata = disegna_punti_critici(image_bytes, punti_totali, abilita_marker=True)
                img_stream = io.BytesIO()
                img_taggata.save(img_stream, format='JPEG', quality=95)
                img_stream.seek(0)
                
                # Inserimento immagine
                row_c.cells[0].add_paragraph().add_run().add_picture(img_stream, width=Inches(2.0))
            else:
                row_c.cells[0].add_paragraph("Immagine non trovata")

            # Testo verbale
            testo = st.session_state.edits.get(f"edit_testo_{idx}", data["trascrizione"])
            paragrafo = row_c.cells[1].add_paragraph(testo)
            run = paragrafo.runs[0]
            run.font.name = font_name
            run.font.size = Pt(font_size)
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            for cell in row_t.cells:
                remove_internal_borders(cell, bottom=True) # Rimuove la riga tra Titolo e Contenuto

            for cell in row_c.cells:
                remove_internal_borders(cell, top=True) # Rimuove la riga tra Titolo e Contenuto
                
            
            # --- PUNTI CRITICI: NUMERAZIONE E SPAZIATURA ---
            for i, p in enumerate(punti_totali, start=1):

                key_punto = f"edit_punto_{idx}_{i}"
                testo_da_scrivere = st.session_state.edits.get(key_punto, p.get('commento', ''))
                testo_finale = f"{i}. {testo_da_scrivere}"
                
                # Creiamo il paragrafo
                p_punto = row_c.cells[1].add_paragraph(testo_finale)
                
                # 1. Rientro (simulazione TAB)
                p_punto.paragraph_format.left_indent = Inches(0.2)
                
                # 2. Spaziatura tra i punti
                p_punto.paragraph_format.space_before = Pt(font_size) 
                
                # 3. Spazio extra per il primo elemento
                if i == 1:
                    p_punto.paragraph_format.space_before = Pt(24)
                
                # Applichiamo font e dimensione
                run_p = p_punto.runs[0]
                run_p.font.name = font_name
                run_p.font.size = Pt(font_size)
            
        # 4. Spostamento XML e pulizia
        table_element = table._tbl
        p_element.addprevious(table_element)
        p_element.getparent().remove(p_element)


    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()