import base64
import json
import io
import time
import random
import streamlit as st
from openai import OpenAI
from streamlit_mic_recorder import mic_recorder
from PIL import Image, ImageDraw
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.shared import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from streamlit_image_coordinates import streamlit_image_coordinates
from streamlit_browser_storage import SessionStorage as Storage
from streamlit_local_storage import LocalStorage





# Metti questa istanza fuori, come variabile globale del modulo
_storage_cache = {}

def get_ls(chiave):
    # Usiamo un dizionario locale al modulo, NON nel session_state 
    # per evitare che il componente venga ricreato durante i rerun
    global _storage_cache
    
    if chiave not in _storage_cache:
        # Questa riga viene eseguita SOLO UNA VOLTA per chiave
        _storage_cache[chiave] = LocalStorage(key=chiave)
        
    return _storage_cache[chiave]

# 2. Funzione di reset
def resetta_tutto_il_sistema():
    st.session_state.anagrafica = {}
    st.session_state.storico_report = []
    st.session_state.edits = {}
    st.session_state.user_data = None
    
    master = get_ls("MASTER_POINTER")
    chiave = master.getItem("chiave_valida")
    
    if chiave:
        try:
            get_ls(chiave).deleteItem("imprendo_dati")
        except: pass
        master.deleteItem("chiave_valida")
            
    st.session_state.ls_registry = {} # Pulisce il registro
    st.toast("Sistema resettato!", icon="🔄")

# 3. Salvataggio
def salva_stato_completo():
    # USA IL GET_LS: NON istanziare LocalStorage direttamente
    master = get_ls("MASTER_POINTER")
    chiave_attuale = master.getItem("chiave_valida")
    
    if not chiave_attuale:
        chiave_attuale = f"storage_{random.randint(10000, 99999)}"
        master.setItem("chiave_valida", chiave_attuale)
    
    localS = get_ls(chiave_attuale)
    
    # [TUA CONVERSIONE BYTES INTATTA COME VOLEVI]
    storico_salvabile = []
    for item in st.session_state.storico_report:
        item_copy = item.copy()
        if "bytes" in item_copy and isinstance(item_copy["bytes"], bytes):
            item_copy["bytes"] = base64.b64encode(item_copy["bytes"]).decode('utf-8')
        storico_salvabile.append(item_copy)

    data = {
        "anagrafica": st.session_state.anagrafica,
        "storico_report": storico_salvabile,
        "edits": st.session_state.edits
    }
    localS.setItem("imprendo_dati", data)
    
    # --- DEBUG DI VERIFICA POST-SCRITTURA ---
    dati_appena_scritti = localS.getItem("imprendo_dati")
    if dati_appena_scritti:
        mandataria_scritta = dati_appena_scritti.get("anagrafica", {}).get("mandataria")
        st.write(f"✅ DEBUG_VERIFICA: Valore letto dal LocalStorage DOPO il salvataggio: '{mandataria_scritta}'")
        if mandataria_scritta != st.session_state.anagrafica['mandataria']:
            st.error("⚠️ ATTENZIONE: Il dato letto è diverso da quello salvato!")
    else:
        st.error("❌ ERRORE: Non riesco a leggere 'imprendo_dati' appena salvato!")

        

# 4. Recupero
def recupera_stato_completo():
    # --- DEBUG DI INIZIO ---
    #st.write("DEBUG RECUPERO: Inizio funzione")
    
    master = LocalStorage(key="MASTER_POINTER")
    chiave_reale = master.getItem("chiave_valida")
    
    st.sidebar.write(f"DEBUG RECUPERO: Chiave trovata: {chiave_reale}")
    
    if not chiave_reale:
        return False
        
    localS = LocalStorage(key=chiave_reale)
    dati = localS.getItem("imprendo_dati")
    
    if dati:
        # --- DEBUG DATI LETTI ---
        #st.write(f"DEBUG RECUPERO: Dati letti dal disco: {list(dati.keys())}")
        #st.write(f"DEBUG RECUPERO: Anagrafica nel disco: {dati.get('anagrafica', {})}")
        
        st.session_state.anagrafica = dati.get("anagrafica", {})
        st.session_state.edits = dati.get("edits", {})
        
        # RIPRISTINO BASE64 -> BYTES
        storico_recuperato = []
        for item in dati.get("storico_report", []):
            item_copy = item.copy()
            if "bytes" in item_copy and isinstance(item_copy["bytes"], str):
                item_copy["bytes"] = base64.b64decode(item_copy["bytes"])
            storico_recuperato.append(item_copy)
            
        st.session_state.storico_report = storico_recuperato
        
        # --- DEBUG FINALE ---
        #st.write("DEBUG RECUPERO: Assegnazione completata.")
        #st.write(f"DEBUG RECUPERO: Session State Anagrafica finale: {st.session_state.anagrafica}")
        
        return True
    
   # st.write("DEBUG RECUPERO: Nessun dato trovato (dati è nullo)")
    return False



def log_sidebar_debug_completo():
    st.sidebar.markdown("---")
    st.sidebar.subheader("🔍 DEBUG LOCAL STORAGE")
    
    master = get_ls("MASTER_POINTER")
    chiave = master.getItem("chiave_valida")
    st.sidebar.write(f"Chiave attiva: **{chiave}**")
    
    if chiave:
        localS = get_ls(chiave)
        dati = localS.getItem("imprendo_dati")
        
        if dati:
            # Iteriamo su tutte le chiavi principali presenti nel JSON salvato
            for sezione, contenuto in dati.items():
                if isinstance(contenuto, list):
                    st.sidebar.write(f"📂 **{sezione}**: {len(contenuto)} elementi")
                elif isinstance(contenuto, dict):
                    st.sidebar.write(f"📂 **{sezione}**: {len(contenuto)} chiavi")
                else:
                    st.sidebar.write(f"📂 **{sezione}**: Trovato")
                
                # Se vuoi vedere il contenuto nel dettaglio (cliccabile)
                with st.sidebar.expander(f"Dettaglio {sezione}"):
                    st.json(contenuto)
        else:
            st.sidebar.error("❌ 'imprendo_dati' è NULL nel LocalStorage!")
    else:
        st.sidebar.warning("⚠️ MASTER_POINTER vuoto.")
    st.sidebar.markdown("---")






def login():
    if "user_data" not in st.session_state:
        st.session_state.user_data = None

    if st.session_state.user_data:
        return st.session_state.user_data

    st.title("🔒 Imprendo")
    st.write("### Il tuo Assistente AI per la sicurezza nei cantieri")
    
    username = st.text_input(
        "Username (Nome)", 
        key="login_username", 
        autocomplete="username"
    ).lower().strip()
    
    password = st.text_input(
        "Password", 
        type="password", 
        key="login_password", 
        autocomplete="current-password"
    )
    
    if st.button("Accedi", use_container_width=True):
        if "utenti" in st.secrets and username in st.secrets["utenti"]:
            db_user = st.secrets["utenti"][username]
            if password == db_user["password"]:
                real_name = db_user.get("nome", username.capitalize())
                
                # 1. Impostiamo i dati utente
                st.session_state.user_data = {
                    "username": username, 
                    "email": db_user["email"], 
                    "nome": real_name,
                    "id": db_user.get("id", "") 
                }
                
                # 2. Inizializziamo l'anagrafica se non esiste
                # if "anagrafica" not in st.session_state:
                #     st.session_state.anagrafica = {}
                
                # ORA recupera i dati: questo andrà a riempire o sovrascrivere 
                # i dizionari vuoti con quelli salvati nel browser
                recupera_stato_completo()
                
                # 4. Ricarichiamo per entrare nell'app
                st.rerun()
            else:
                st.error(f"❌ **Password errata!**")
        else:
            st.error(f"❌ **Utente non trovato!**")
    return None





def ottieni_account_exchange(user_email):
    from O365 import Account
    

    config = st.secrets["microsoft_exchange"]
    if not config:
        st.error(f"⚠️ Configurazione non trovata per il dominio: {dominio}")
        return None
        
    credentials = (config["client_id"], config["client_secret"])
    tenant_id = config["tenant_id"]
    
    account = Account(credentials, auth_flow_type='credentials', tenant_id=tenant_id)
    
    try:
        if account.authenticate():
            return account
        return None
    except Exception as e:
        st.error(f"❌ Errore autenticazione Azure ({dominio}): {e}")
        return None
    

def invia_report_via_email_graph(doc_bytes, nome_file, user_email):
    import requests
    import base64

    try:
        # 1. Recupero credenziali
        config = st.secrets["microsoft_exchange"]
        client_id = config["client_id"]
        client_secret = config["client_secret"]
        tenant_id = config["tenant_id"]

        # 2. Ottenimento Token
        url_oauth = f"https://login.microsoftonline.com/{tenant_id}/oauth2/v2.0/token"
        payload_oauth = {
            "grant_type": "client_credentials",
            "client_id": client_id,
            "client_secret": client_secret,
            "scope": "https://graph.microsoft.com/.default"
        }
        risposta_oauth = requests.post(url_oauth, data=payload_oauth)
        if risposta_oauth.status_code != 200:
            return False, f"Errore Auth: {risposta_oauth.text}"
        
        token = risposta_oauth.json().get("access_token")

        # 3. Preparazione allegato in Base64 (come nel tuo codice collaudato)
        encoded_content = base64.b64encode(doc_bytes).decode('utf-8')
        
        email_payload = {
            "message": {
                "subject": f"Report Sicurezza Cantiere - {time.strftime('%Y-%m-%d %H:%M')}",
                "body": {
                    "contentType": "HTML",
                    "content": "<p>In allegato il report di sicurezza generato.</p>"
                },
                "toRecipients": [{"emailAddress": {"address": user_email}}],
                "attachments": [
                    {
                        "@odata.type": "#microsoft.graph.fileAttachment",
                        "name": nome_file,
                        "contentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        "contentBytes": encoded_content
                    }
                ]
            },
            "saveToSentItems": "true"
        }

        # 4. Invio
        url_api = f"https://graph.microsoft.com/v1.0/users/{user_email}/sendMail"
        headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
        
        risposta = requests.post(url_api, json=email_payload, headers=headers)

        if risposta.status_code == 202:
            return True, "Email inviata con successo!"
        else:
            return False, f"Errore API ({risposta.status_code}): {risposta.text}"

    except Exception as e:
        return False, str(e)
    

def set_global_styles():
    st.markdown(
        """
        <style>
        /* CSS per immagini standard e dentro expander */
        [data-testid="stExpander"] img,
        [data-testid="stImage"] img {
            max-width: 100% !important;
            height: auto !important;
            display: block;
            margin-left: auto;
            margin-right: auto;
        }
        
        /* Forza il contenitore dell'immagine a rispettare la larghezza del padre */
        [data-testid="stExpander"] [data-testid="stImage"],
        [data-testid="stImage"] {
            width: 100% !important;
        }
        </style>
        """,
        unsafe_allow_html=True
    )


def set_bg_color(color, status_text=None):
    # CSS per cambiare lo sfondo e creare il banner fisso
    banner_html = ""
    if status_text:
        banner_html = f"""
        <div style="position: fixed; top: 0; left: 0; width: 100%; background-color: #333; 
                    color: white; text-align: center; padding: 10px; z-index: 9999; font-weight: bold;">
            {status_text}
        </div>
        """
    
    st.markdown(
        f"""
        <style>
        .stApp {{ background-color: {color} !important; }}
        </style>
        {banner_html}
        """,
        unsafe_allow_html=True
    )




def disegna_punti_critici(image_bytes, lista_punti, abilita_marker=True):
    img = Image.open(io.BytesIO(image_bytes)).convert("RGB")

    # Se i marker sono disabilitati salta la funzione
    if not abilita_marker:
        return img
    
    else:
        draw = ImageDraw.Draw(img)
        width, height = img.size
        #colori = ["red", "blue", "yellow", "cyan", "magenta", "lime", "orange", "purple"]
        colore = "red"
        
        for i, punto in enumerate(lista_punti):
            #colore = colori[i % len(colori)]
            coord = punto.get('coordinate', {'x': 50, 'y': 50})
            x_raw = coord.get('x')
            y_raw = coord.get('y')

            # --- FIX: Se x o y sono None, salta questo punto ---
            if x_raw is None or y_raw is None:
                continue

            size = 30
            
            # --- CORREZIONE: Normalizzazione dinamica ---
            x_raw = coord.get('x', 40)
            y_raw = coord.get('y', 50)
            x = (x_raw / 1000) * width if x_raw > 100 else (x_raw / 100) * width
            y = (y_raw / 1000) * height if y_raw > 100 else (y_raw / 100) * height
            
            # Disegna cerchio e testo
            draw.ellipse([x-size, y-size, x+size, y+size], outline=colore, width=3)
            draw.text((x, y), str(i + 1), fill=colore, font_size=40, anchor="mm")
            
        return img


def analizza_sicurezza_cantiere(audio_bytes, image_file):
    """
    Analizza una singola immagine e le note vocali.
    image_file: l'oggetto file di streamlit (quello che ha .name e .getvalue())
    """
    client = OpenAI(api_key=st.secrets["openai_key"])
    
    # 1. Trascrizione vocale
    audio_file = io.BytesIO(audio_bytes)
    audio_file.name = "audio.wav"
    transcript = client.audio.transcriptions.create(model="whisper-1", file=audio_file)
    
    # PROMPT ORIGINALE (Mantenuto identico)
    system_prompt = """
    Sei un Ispettore Tecnico della Sicurezza (D.Lgs 81/08).
    
    Tuo compito:
    1. Analizza l'immagine e le note vocali dell'ispettore.
    2. ELABORAZIONE DISCORSO: Trasforma la trascrizione vocale dell'ispettore in un "Verbale Tecnico Formale" sintetico, professionale e preciso (rimuovi intercalari, ripetizioni o linguaggio colloquiale).
    3. Identifica i punti critici dell'audio e mappali su coordinate in pixel (x,y).
    4. Identifica autonomamente ulteriori criticità (segnali come 'AI Integrativa').
    5. Determina l'oggetto delle analisi in circa 5 parole.
    6. Il "riassunto_generale" deve essere di 10 parole massimo.
    
    Restituisci JSON: 
    {
        "verbale_tecnico": "Il testo rielaborato e formale basato sulle note vocali",
        "riassunto_generale": "...",
        "punti_critici": [
            {
                "elemento": "...",
                "origine": "Ispettore" o "AI Integrativa",
                "commento": "...",
                "rischio": "...",
                "urgenza": "...",
                "coordinate": {"x": 50, "y": 50},
                "oggetto": "..."
            }
        ]
    }
    """

    
    # 2. Codifica immagine
    b64 = base64.b64encode(image_file.getvalue()).decode('utf-8')
    
    payload = [
        {"type": "text", "text": f"Note vocali ispettore: {transcript.text}"},
        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}}
    ]
    
    # 3. Chiamata API
    resp = client.chat.completions.create(
        model="gpt-4o", 
        messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": payload}], 
        response_format={"type": "json_object"}
    )
    
    dati = json.loads(resp.choices[0].message.content)
    
    # Usiamo il "verbale_tecnico" generato dall'AI come testo finale
    testo_formale = dati.get("verbale_tecnico", transcript.text) 
    
    report_completo = {
        "riassunto_generale": dati.get("riassunto_generale", ""),
        "analisi_per_immagine": [{
            "nome_file": image_file.name, 
            "punti_critici": dati.get("punti_critici", [])
        }]
    }
        
    return report_completo, testo_formale



def integra_sicurezza_cantiere(audio_bytes, image_file, verbale_attuale):
    """
    Integra un verbale esistente con nuove note vocali.
    image_file: l'oggetto file (o MockFile) con .name e .getvalue()
    verbale_attuale: stringa contenente il verbale già editato dall'utente
    """
    client = OpenAI(api_key=st.secrets["openai_key"])
    
    # 1. Trascrizione della nuova integrazione
    audio_file = io.BytesIO(audio_bytes)
    audio_file.name = "audio.wav"
    transcript = client.audio.transcriptions.create(model="whisper-1", file=audio_file)
    
    # 2. Prompt per l'integrazione
    system_prompt = f"""
    Sei un Ispettore Tecnico della Sicurezza (D.Lgs 81/08).
    
    Hai un verbale attuale:
    "{verbale_attuale}"
    
    L'ispettore ha aggiunto nuove note vocali:
    "{transcript.text}"
    
    Tuo compito:
    1. Integra le nuove informazioni nel verbale tecnico formale esistente, mantenendo lo stile professionale.
    2. Aggiorna la lista dei punti critici includendo quelli nuovi e mantenendo quelli vecchi validi.
    3. Se le nuove note rendono obsoleti alcuni punti vecchi, correggili o rimuovili.
    4. Il "riassunto_generale" deve rimanere di 10 parole massimo.
    5. Il JSON deve contenere TUTTI i punti critici aggiornati e completi.
    
    Restituisci JSON con la stessa struttura precedente: 
    {{
        "verbale_tecnico": "Il verbale aggiornato e completo",
        "riassunto_generale": "...",
        "punti_critici": [
            {{
                "elemento": "...",
                "origine": "...",
                "commento": "...",
                "rischio": "...",
                "urgenza": "...",
                "coordinate": {{"x": 50, "y": 50}},
                "oggetto": "..."
            }}
        ]
    }}
    """
    
    # 3. Codifica immagine
    b64 = base64.b64encode(image_file.getvalue()).decode('utf-8')
    
    payload = [
        {"type": "text", "text": "Aggiorna il verbale in base alle note integrative."},
        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}}
    ]
    
    # 4. Chiamata API
    resp = client.chat.completions.create(
        model="gpt-4o", 
        messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": payload}], 
        response_format={"type": "json_object"}
    )
    
    dati = json.loads(resp.choices[0].message.content)
    
    # Ritorno del report aggiornato
    report_completo = {
        "riassunto_generale": dati.get("riassunto_generale", ""),
        "analisi_per_immagine": [{
            "nome_file": image_file.name, 
            "punti_critici": dati.get("punti_critici", [])
        }]
    }
    
    return report_completo, dati.get("verbale_tecnico", verbale_attuale)



def remove_internal_borders(cell, top=False, bottom=False):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    if top:
        top_el = OxmlElement('w:top')
        top_el.set(qn('w:val'), 'nil')
        tcBorders.append(top_el)
    if bottom:
        bottom_el = OxmlElement('w:bottom')
        bottom_el.set(qn('w:val'), 'nil')
        tcBorders.append(bottom_el)




def genera_report_finale(storico):


    doc = Document("Template.docx")

    font_name = st.session_state.settings.get("font", "Arial")
    font_size = st.session_state.settings.get("size", 9)

    # --- FUNZIONE DI SOSTITUZIONE ROBUSTA ---
    def replace_text_in_doc(doc, key, value, font_name, font_size):

        placeholder = f"{{{{{key}}}}}"

        # Funzione interna per applicare lo stile al paragrafo
        def apply_style_to_para(para):
            for run in para.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)

        # Elaborazione paragrafi
        for para in doc.paragraphs:
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, value)
                apply_style_to_para(para)
                
        # Elaborazione tabelle
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if placeholder in para.text:
                            para.text = para.text.replace(placeholder, value)
                            apply_style_to_para(para)

    # Eseguiamo la sostituzione per ogni chiave
    if "anagrafica" in st.session_state:
        for key, value in st.session_state.anagrafica.items():
            replace_text_in_doc(doc, key, value, font_name, font_size)
    
    # 1. Trova il paragrafo segnaposto
    target_paragraph = None
    for p in doc.paragraphs:
        if "###TABELLA_ANALISI###" in p.text:
            target_paragraph = p
            break
            
    if target_paragraph:
        p_element = target_paragraph._element
        
        # 2. Crea la tabella
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        
        # Intestazione
        hdr = table.add_row().cells
        hdr[0].text = "ANALISI TECNICHE"

        for cell in hdr:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), 'D9D9D9') # Inserisci qui il codice esadecimale del colore
            tcPr.append(shd)

    
        # 3. Ciclo sullo storico
        for idx, data in enumerate(storico):
            # Recupera i punti aggiornati (inclusi quelli manipolati dall'utente)
            punti_totali = [p for img_data in data["report"].get("analisi_per_immagine", []) for p in img_data['punti_critici']]
            
            # Riga titolo con riassunto specifico
            riassunto_specifico = data["report"].get("riassunto_generale", "Analisi Tecnica")
            
            # Riga titolo
            row_t = table.add_row()
            for cell in row_t.cells:
                # 1. Imposta l'allineamento verticale
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
                # 2. Rimuovi margini interni del paragrafo che potrebbero dare l'illusione di decentramento
                for para in cell.paragraphs:
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
            
            # Formattazione FOTOGRAFIA X
            paragrafo_foto = row_t.cells[0].add_paragraph(f"FOTOGRAFIA {idx + 1}")
            run_foto = paragrafo_foto.runs[0]
            run_foto.font.name = font_name
            run_foto.font.size = Pt(font_size)
            run_foto.bold = True
            
            # Formattazione RIASSUNTO
            paragrafo_t = row_t.cells[1].add_paragraph(riassunto_specifico)
            run_t = paragrafo_t.runs[0]
            run_t.font.name = font_name
            run_t.font.size = Pt(font_size)
            run_t.bold = True
            
            
            # Riga contenuto
            row_c = table.add_row()
            
            # --- DISEGNO IMMAGINE CON TAG ---
            img_taggata = disegna_punti_critici(data["bytes"], punti_totali, abilita_marker=True)
            img_stream = io.BytesIO()
            img_taggata.save(img_stream, format='JPEG', quality=95)
            img_stream.seek(0)
            
            # Inserimento immagine
            row_c.cells[0].add_paragraph().add_run().add_picture(img_stream, width=Inches(2.0))

            # Testo verbale
            testo = st.session_state.edits.get(f"edit_testo_{idx}", data["trascrizione"])
            paragrafo = row_c.cells[1].add_paragraph(testo)
            run = paragrafo.runs[0]
            run.font.name = font_name
            run.font.size = Pt(font_size)
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            for cell in row_t.cells:
                remove_internal_borders(cell, bottom=True) # Rimuove la riga tra Titolo e Contenuto

            for cell in row_c.cells:
                remove_internal_borders(cell, top=True) # Rimuove la riga tra Titolo e Contenuto
                
            
            # --- PUNTI CRITICI: NUMERAZIONE E SPAZIATURA ---
            for i, p in enumerate(punti_totali, start=1):

                key_punto = f"edit_punto_{idx}_{i}"
                testo_da_scrivere = st.session_state.edits.get(key_punto, p.get('commento', ''))
                testo_finale = f"{i}. {testo_da_scrivere}"
                
                # Creiamo il paragrafo
                p_punto = row_c.cells[1].add_paragraph(testo_finale)
                
                # 1. Rientro (simulazione TAB)
                p_punto.paragraph_format.left_indent = Inches(0.2)
                
                # 2. Spaziatura tra i punti
                p_punto.paragraph_format.space_before = Pt(font_size) 
                
                # 3. Spazio extra per il primo elemento
                if i == 1:
                    p_punto.paragraph_format.space_before = Pt(24)
                
                # Applichiamo font e dimensione
                run_p = p_punto.runs[0]
                run_p.font.name = font_name
                run_p.font.size = Pt(font_size)
            
        # 4. Spostamento XML e pulizia
        table_element = table._tbl
        p_element.addprevious(table_element)
        p_element.getparent().remove(p_element)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()



def get_img_bytes(pil_img):
    img_byte_arr = io.BytesIO()
    # Salviamo in formato JPEG con alta qualità
    pil_img.save(img_byte_arr, format='JPEG', quality=95)
    return img_byte_arr.getvalue()


def trascrivi_in_campo(campo_anagrafica, audio_bytes):
    """Trascrive l'audio e aggiorna il valore nel session_state."""
    client = OpenAI(api_key=st.secrets["openai_key"])
    audio_file = io.BytesIO(audio_bytes)
    audio_file.name = "audio.wav"
    
    # Trascrizione
    transcript = client.audio.transcriptions.create(model="whisper-1", file=audio_file)
    
    # Aggiornamento dello stato (concatenazione)
    testo_attuale = st.session_state.anagrafica.get(campo_anagrafica, "")
    if testo_attuale:
        st.session_state.anagrafica[campo_anagrafica] = testo_attuale + "\n" + transcript.text
    else:
        st.session_state.anagrafica[campo_anagrafica] = transcript.text

def campo_con_audio(label, key_campo, help_text="", tipo="area"):
    c1, c2 = st.columns([0.85, 0.15])
    
    with c1:
        if tipo == "area":
            valore = st.text_area(label, value=st.session_state.anagrafica.get(key_campo, ""), help=help_text)
        else:
            valore = st.text_input(label, value=st.session_state.anagrafica.get(key_campo, ""), help=help_text)
        st.session_state.anagrafica[key_campo] = valore
        
    with c2:
        st.write("###")
        # Aggiungiamo un hash per distinguere l'audio
        audio = mic_recorder(start_prompt="🎤", stop_prompt="⏹️", key=f"rec_{key_campo}")
        
        if audio:
            # Creiamo una chiave unica per questo specifico audio basata sul contenuto
            audio_id = hash(str(audio['bytes']))
            
            # Controlliamo se abbiamo già elaborato questo esatto audio
            if st.session_state.get(f"last_audio_{key_campo}") != audio_id:
                with st.spinner("Trascrizione..."):
                    trascrivi_in_campo(key_campo, audio['bytes'])
                    # Segniamo l'audio come "già elaborato"
                    st.session_state[f"last_audio_{key_campo}"] = audio_id
                    st.rerun()



def elabora_anagrafica_ai(audio_bytes):
    client = OpenAI(api_key=st.secrets["openai_key"])
    audio_file = io.BytesIO(audio_bytes)
    audio_file.name = "audio.wav"
    
    # 1. Trascrizione
    transcript = client.audio.transcriptions.create(model="whisper-1", file=audio_file)
    
    # 2. Estrazione dati strutturati con AI
    prompt = f"""
    Sei l'assistente di un ispettore sicurezza nei cantieri che sta dirigendo un report.
    Estrai le informazioni dal seguente testo e restituiscile in formato JSON:
    "{transcript.text}"
    
    Restituisci JSON con queste chiavi (se un dato manca, metti una stringa vuota):
    {{
        "mandataria": "Elenco mandatarie",
        "mandante": "Elenco mandanti",
        "committente": "Nome committente",
        "indirizzo": "Indirizzo committente",
        "città": "Città",
        "provincia": "Provincia in formato (XX)",
        "commessa": "Commessa del committente espressa in modo formale ed esaustiva",
        "oggetto": "Dettaglio del lavoro commissionato espresso in modo formale ed esaustivo"
    }}
    """
    
    resp = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        response_format={"type": "json_object"}
    )
    
    return json.loads(resp.choices[0].message.content)


def elabora_campo_tecnico_ai(audio_bytes, nome_campo):
    client = OpenAI(api_key=st.secrets["openai_key"])
    audio_file = io.BytesIO(audio_bytes)
    audio_file.name = "audio.wav"
    transcript = client.audio.transcriptions.create(model="whisper-1", file=audio_file)
    
    #  1. Attività: Descrizione delle attività lavorative in corso nel cantiere, usa una terminologia tecnica (D.Lgs 81/08).
    # 2. Coordinamento: Descrizione delle azioni di coordinamento e vigilanza svolte dall'ispettore.
    # 3. Personale: Elenco delle figure professionali presenti, indicazioni sull'autorizzazione dei lavoratori di accedere al cantiere ed elenco delle eventuali ditte subappaltatrici presenti.
    # 4. Verbale: Elenco delle prescrizioni, delle sospensioni o di altri verbali rilasciati durante il sopralluogo.

    prompt = f"""
    Sei l'assistente di un ispettore sulla sicurezza nei cantieri e stai ascoltando il suo resoconto: "{transcript.text}".
    Estrai le informazioni in modo formale e tecnico per il report di fine ispezione.

    REGOLE
    1. Se la trascrizione è assente o inutilizzabile, restituisci il campo vuoto.
    2. Non scrivere frasi introduttive quali: "il resconto dell'ispettore...", "L'ispettore ha indicato..."
    """
    
    resp = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}]
    )
    return resp.choices[0].message.content


# APP PRINCIPALE
recupera_stato_completo() 

# --- INIZIALIZZAZIONE FORZATA (DEVE ESSERE IN CIMA) ---
if "anagrafica" not in st.session_state:
    st.session_state.anagrafica = {}
    recupera_stato_completo() 

# Assicuriamoci che esistano le chiavi base per evitare KeyError
chiavi_base = ["mandataria", "mandante", "committente", "indirizzo", "città", "provincia", 
               "commessa", "oggetto", "attività", "coordinamento", "personale", "verbali"]
for k in chiavi_base:
    if k not in st.session_state.anagrafica:
        st.session_state.anagrafica[k] = ""

if "anagrafica_version" not in st.session_state:
    st.session_state.anagrafica_version = 0


# --- ORA CHIAMA IL LOGIN ---
utente_connesso = login()

set_global_styles()

log_sidebar_debug_completo()

if "app_state" not in st.session_state:
    status_msg = None
    color = "white"
    st.session_state.app_state = "ready"
if st.session_state.app_state == "ready":
    set_bg_color("#ffffff") 
if st.session_state.app_state == "working":
    color = "#D0AD00"
    status_msg = "⚠️ ANALISI IN CORSO - NON INTERAGIRE"
elif st.session_state.app_state == "done":
    color = "#89D889"
    status_msg = "✅ ANALISI COMPLETATA CON SUCCESSO"
else:
    status_msg = None
    color = "white"


set_bg_color(color, status_msg)

# --- 3. CONTENUTO PRINCIPALE ---
if utente_connesso:

    status_placeholder = st.empty()
    
    if st.sidebar.button("💾 SALVA BOZZA", type="primary"):
        status_placeholder.warning("Salvataggio...")
        
        campi_anagrafica = ["mandataria", "mandante", "committente", "indirizzo", "città", "provincia"]
        for c in campi_anagrafica:
            key_widget = f"widget_{c}_{st.session_state.get('anagrafica_version', 0)}"
            if key_widget in st.session_state:
                st.session_state.anagrafica[c] = st.session_state[key_widget]
                
        # Esempio per Commessa/Oggetto
        campi_tecnici = ["commessa", "oggetto", "attività", "coordinamento", "personale", "verbali"]
        for c in campi_tecnici:
            key_widget = f"widget_{c}_{st.session_state.get('anagrafica_version', 0)}"
            if key_widget in st.session_state:
                st.session_state.anagrafica[c] = st.session_state[key_widget]

        # 2. Ora salviamo il session_state (che è stato appena sincronizzato)
        salva_stato_completo()
        st.sidebar.success("Dati aggiornati e salvati!")

        time.sleep(1)
        st.rerun()

    st.sidebar.subheader("Reset App")
    if st.sidebar.button("🔄 Inizia da zero"):
        status_placeholder.warning("Reset totale in corso...")
        resetta_tutto_il_sistema()
        time.sleep(1)
        st.rerun()
    st.sidebar.divider()


    # IMPOSTAZIONI 
    st.sidebar.header("⚙️ Impostazioni")

    # Mostra marker
    mostra_marker = st.sidebar.toggle("Mostra Marker sulla foto", value=True)

    # Scelta font
    font_s = st.sidebar.selectbox("Font", ["Arial", "Calibri", "Times New Roman"], index=0)
    
    # Scelta dimensione
    size_s = st.sidebar.slider("Dimensione Font (pt)", min_value=7, max_value=16, value=9)
    
    # Salviamo nello stato
    st.session_state.settings = {
        "font": font_s,
        "size": size_s
    }

    # Inizializzazione variabili di stato
    if "storico_report" not in st.session_state: st.session_state.storico_report = []
    if "edits" not in st.session_state: st.session_state.edits = {}

    tab1, tab2, tab3 = st.tabs(["🚀 Caricamenti", "📋 Analisi Tecniche", "👤 Report"])

    # --- TAB 1: ACQUISIZIONE E ANALISI ---
    with tab1:
        st.subheader("📸 Carica e descrivi")
        file = st.file_uploader("Carica una foto", type=["jpg", "png", "jpeg"], key="uploader_live")
            
        audio = mic_recorder(
            start_prompt="🟢 AVVIA REGISTRAZIONE", 
            stop_prompt="🛑 FERMA REGISTRAZIONE E AVVIA ANALISI", 
            key='recorder_live'
        )

        # Visualizzazione immagine di guida
        if file is not None:
            st.image(file, caption="Immagine di riferimento per il sopralluogo", use_container_width=True)
            set_bg_color("#ffffff") 

        if audio and file:

            audio_hash = hash(str(audio['bytes']))
            
            # Evita esecuzioni multiple
            if st.session_state.get("last_audio_hash") != audio_hash:

                with st.spinner("Analisi in corso..."):
                    # Esecuzione Analisi
                    set_bg_color("#D0AD00")
                    st.session_state.app_state = "working"
                    report, testo = analizza_sicurezza_cantiere(audio['bytes'], file)
                    
                    # Aggiornamento storico
                    if "storico_report" not in st.session_state: 
                        st.session_state.storico_report = []
                    
                    st.session_state.storico_report.append({
                        "nome_file": file.name, 
                        "report": report, 
                        "trascrizione": testo, 
                        "bytes": file.getvalue(), 
                        "version": 1
                    })
                    
                    st.session_state.last_audio_hash = audio_hash
                    salva_stato_completo()
                    st.session_state.app_state = "done"
                    set_bg_color("#b3ff99")
                    time.sleep(2)
                    st.rerun()

    # --- TAB 2: VISUALIZZAZIONE E GESTIONE (REWORK/INTEGRAZIONE) ---
    with tab2:
        if st.session_state.storico_report:
            for idx, data in enumerate(st.session_state.storico_report):
                expander_key = f"expander_idx_{idx}_id_{hash(str(data))}"

                nome_file = data["nome_file"]
                report = data["report"]
                ver = data.get("version", 1) 
                punti_totali = [p for img_data in report.get("analisi_per_immagine", []) for p in img_data['punti_critici']]
                titolo = report.get("riassunto_generale", f"Analisi {nome_file}")
                
                # Usa l'indice, che è costante finché l'elemento non viene rimosso
                with st.expander(f"🔍 {titolo.upper()} ({nome_file})", expanded=True, key=f"expander_{idx}"):
                    col1, col2 = st.columns([1, 1])
                    with col1:
                        
                        # 2. IMMAGINE INTERATTIVA
                        img_da_disegnare = data.get("bytes")
                        if img_da_disegnare:
                            # Caso standard: hai i bytes, disegna tutto
                            img_display = disegna_punti_critici(img_da_disegnare, punti_totali, abilita_marker=mostra_marker)
                        else:
                            # Caso fallback: non hai i bytes (sono stati filtrati), 
                            # ma forse hai il file caricato nel widget uploader o altro riferimento?
                            # Se non c'è, carica un'immagine neutra per evitare il crash
                            
                            img_display = Image.new('RGB', (300, 300), color=(200, 200, 200))
                            st.warning("Immagine originale non caricata in memoria.")


                        # 2. Rendiamo l'immagine cliccabile (al posto di st.image)
                        # Nota: dobbiamo passare l'immagine PIL (img_display)
                        click_data = streamlit_image_coordinates(
                            img_display, 
                            key=f"img_click_{idx}",
                            width=350
                        )

                        # 3. Logica per aggiungere il punto se l'utente clicca
                        if click_data is not None:
                            if st.button("📍 Fissa punto qui", key=f"fix_{idx}_{ver}"):
                                w, h = img_display.size
                                x_norm = (click_data['x'] / w) * 1000
                                y_norm = (click_data['y'] / h) * 1000
                                
                                # CERCHIAMO IL PRIMO PUNTO CON COORDINATE NONE
                                punto_vuoto_trovato = False
                                for p in report['analisi_per_immagine'][0]['punti_critici']:
                                    if p.get('coordinate', {}).get('x') is None:
                                        p['coordinate'] = {'x': x_norm, 'y': y_norm}
                                        punto_vuoto_trovato = True
                                        break
                                
                                # SE NON CI SONO PUNTI VUOTI, NE AGGIUNGIAMO UNO NUOVO
                                if not punto_vuoto_trovato:
                                    report['analisi_per_immagine'][0]['punti_critici'].append({
                                        "elemento": "Punto Manuale",
                                        "commento": "Aggiunto manualmente",
                                        "coordinate": {"x": x_norm, "y": y_norm},
                                        "oggetto": "Nota manuale"
                                    })
                                
                                st.session_state.storico_report[idx]['report'] = report
                                salva_stato_completo()
                                st.rerun()

                        
                        # 1. PULSANTI DI SISTEMA (Spostati in alto per evitare conflitti)
                        c1, c2, c3, c4 = st.columns(4)

                        with c1:
                            if st.button("🗑️ Elimina", key=f"del_button_{idx}"):
                                # 1. Rimuovi i dati
                                del st.session_state.storico_report[idx]
                                
                                # 2. Pulisci le chiavi dei widget associati a quell'indice
                                chiavi_da_rimuovere = [k for k in st.session_state.keys() if f"_{idx}" in k]
                                for k in chiavi_da_rimuovere:
                                    del st.session_state[k]
                                    
                                salva_stato_completo()
                                st.rerun()

                        # if c2.button("🔄 Rifai", key=f"redo_{idx}"):
                        #     st.session_state.active_recorder = {"idx": idx, "mode": "rework"}
                        #     salva_stato_completo()
                        #     st.rerun()

                        # if c3.button("➕ Integra", key=f"int_{idx}"):
                        #     st.session_state.active_recorder = {"idx": idx, "mode": "integration"}
                        #     salva_stato_completo()
                        #     st.rerun()
                        
                        # NUOVO PULSANTE: Svuota solo le coordinate (i testi restano!)

                        if c4.button("🧹 Svuota Marker", key=f"clear_markers_{idx}"):
                            for img_data in report.get("analisi_per_immagine", []):
                                for p in img_data['punti_critici']:
                                    p['coordinate'] = {'x': None, 'y': None} # Rende i cerchi invisibili
                            st.session_state.storico_report[idx]['report'] = report
                            salva_stato_completo()
                            st.rerun()

                        # --- REGISTRATORE CONTESTUALE ---
                        if st.session_state.get("active_recorder") and st.session_state.active_recorder["idx"] == idx:
                            st.info(f"🎤 Registra per {st.session_state.active_recorder['mode'].upper()}...")
                            audio_attivo = mic_recorder(start_prompt="⏺️ AVVIA", stop_prompt="⏹️ ANALIZZA", key=f"rec_{idx}")
                            
                            if audio_attivo:
                                with st.spinner("Analisi in corso..."):
                                    class MockFile:
                                        def __init__(self, name, data): self.name = name; self._data = data
                                        def getvalue(self): return self._data
                                    
                                    file_mock = MockFile(data["nome_file"], data["bytes"])
                                    mode = st.session_state.active_recorder["mode"]
                                    
                                    if mode == "rework":
                                        report, testo = analizza_sicurezza_cantiere(audio_attivo['bytes'], file_mock)
                                    else:
                                        verbale_attuale = st.session_state.edits.get(f"edit_testo_{idx}", data["trascrizione"])
                                        report, testo = integra_sicurezza_cantiere(audio_attivo['bytes'], file_mock, verbale_attuale)
                                    
                                    new_v = st.session_state.get("version_counter", 0) + 1
                                    st.session_state.version_counter = new_v
                                    st.session_state.storico_report[idx].update({"report": report, "trascrizione": testo, "version": new_v})
                                    salva_stato_completo()
                                    
                                    # Reset Edits
                                    for k in [k for k in st.session_state.edits.keys() if f"_{idx}" in k]: del st.session_state.edits[k]
                                    st.session_state.edits[f"edit_testo_{idx}"] = testo
                                    salva_stato_completo()
                                    
                                    del st.session_state.active_recorder
                                    st.rerun()
                    
                    with col2:
                        st.markdown("#### Analisi")
                        
                        # 1. Inizializzazione dati (solo se mancano)
                        key_testo = f"edit_testo_{idx}"
                        if key_testo not in st.session_state.edits:
                            st.session_state.edits[key_testo] = data["trascrizione"]

                        # 2. Sincronizzazione: Se l'AI ha generato un nuovo testo (es. dopo un rework), 
                        # aggiorniamo il valore SOLO SE l'utente non ha ancora iniziato a modificare manualmente
                        # oppure se forziamo un reset.
                        # Consiglio: non usare st.session_state.edits[key_testo] != data["trascrizione"] 
                        # qui dentro se vuoi che l'utente possa editare senza che l'AI sovrascriva.
                        # Fai l'aggiornamento solo quando chiami la funzione "rework" (nel Tab 2).

                        # 3. WIDGET CON CHIAVE STATICA (La chiave NON DEVE dipendere da ver)
                        valore_attuale = st.session_state.edits[key_testo]

                        st.session_state.edits[key_testo] = st.text_area(
                            "Modifica il verbale:", 
                            value=valore_attuale, 
                            height=230,
                            key=key_testo,  
                            on_change=salva_stato_completo 
                        )
                        
                        st.markdown("#### ⚠️ Punti critici rilevati")
                        
                        # Usiamo un contatore sicuro
                        for idx_p, p in enumerate(punti_totali):
                            # Creiamo un ID unico che non dipende dal numero di elementi nella lista
                            # Se il punto ha un 'id' nel dizionario, usa quello. Altrimenti usa le coordinate.
                            id_univoco = p.get('id', f"x{p.get('coordinate',{}).get('x')}_y{p.get('coordinate',{}).get('y')}_{idx_p}")
                            
                            c_punto1, c_punto2 = st.columns([0.9, 0.1])
                            
                            with c_punto1:
                                key_punto = f"edit_punto_{idx}_{id_univoco}"
                                if key_punto not in st.session_state.edits:
                                    st.session_state.edits[key_punto] = p.get('commento', '')
                                
                                st.session_state.edits[key_punto] = st.text_area(
                                    f"{idx_p + 1}. {p.get('elemento', 'Punto')} ({p.get('oggetto', 'Nota')})",
                                    value=st.session_state.edits[key_punto],
                                    height=130,
                                    key=key_punto,
                                    on_change=salva_stato_completo
                                )
                            
                            with c_punto2:
                                # Il bottone usa la stessa chiave univoca, non i
                                if st.button("❌", key=f"del_punto_{idx}_{id_univoco}"):
                                    # Rimuovi il punto e forza il rerun
                                    for img_data in report.get("analisi_per_immagine", []):
                                        if p in img_data['punti_critici']:
                                            img_data['punti_critici'].remove(p)
                                            st.session_state.storico_report[idx]['report'] = report
                                            salva_stato_completo()
                                            st.rerun()

        else:
            st.info("Esegui un'analisi per vedere i risultati qui.")

    
    with tab3:


        with st.expander("👤 Anagrafiche"):
            with st.container():

                # 1. Pulsante di registrazione unico
                audio_data = mic_recorder(key="rec_anagrafica_totale", start_prompt="🎤", stop_prompt="⏹️")
                
                if audio_data:
                    audio_hash = hash(str(audio_data['bytes']))
                    if st.session_state.get("last_anagrafica_hash") != audio_hash:
                        with st.spinner("L'AI sta estraendo i dati..."):
                            set_bg_color("#D0AD00")

                            dati = elabora_anagrafica_ai(audio_data['bytes'])
                            
                            # Aggiornamento dati
                            st.session_state.anagrafica["mandataria"]   = str(dati.get("mandataria", "")).replace(", ", "\n")
                            st.session_state.anagrafica["mandante"]     = str(dati.get("mandante", "")).replace(", ", "\n")
                            st.session_state.anagrafica.update({
                                "committente": dati.get("committente", ""),
                                "indirizzo": dati.get("indirizzo", ""),
                                "città": dati.get("città", ""),
                                "provincia": dati.get("provincia", "")
                            })

                            # !!! QUESTO È QUELLO CHE MANCAVA !!!
                            # Incrementiamo la versione per cambiare la chiave dei widget
                            st.session_state.anagrafica_version += 1
                            
                            salva_stato_completo()
                            st.session_state.last_anagrafica_hash = audio_hash
                            set_bg_color("#b3ff99")
                            time.sleep(1)
                            st.rerun()
                
                # Lista definita fuori dal loop
                campi = [
                    ("mandataria", "Mandataria/e", "area"),
                    ("mandante", "Mandante/i", "area"),
                    ("committente", "Ragione Sociale Committente", "input"),
                    ("indirizzo", "Indirizzo", "input"),
                    ("città", "Città", "input"),
                    ("provincia", "Provincia", "input")
                ]

                salt = st.session_state.get("anagrafica_version", 0)
                
                for campo_id, label, tipo in campi:
                    
                        key_widget = f"widget_{campo_id}_{salt}"
                        
                        # ASSEGNAZIONE DIRETTA NEL WIDGET
                        if tipo == "area":
                            st.session_state.anagrafica[campo_id] = st.text_area(
                                label, 
                                value=st.session_state.anagrafica.get(campo_id, ""),
                                key=key_widget,
                                on_change=salva_stato_completo
                            )
                        else:
                            st.session_state.anagrafica[campo_id] = st.text_input(
                                label, 
                                value=st.session_state.anagrafica.get(campo_id, ""),
                                key=key_widget,
                                on_change=salva_stato_completo
                            )
                        

        with st.expander("📝 Commessa e Oggetto"):
            with st.container():

                # Recuperiamo il salt corrente (se non esiste, partiamo da 0)
                salt = st.session_state.get("anagrafica_version", 0)

                # Lista di tuple per rendere il codice più compatto e DRY (Don't Repeat Yourself)
                campi_tecnici = [
                    ("commessa", "Commessa", "rec_commessa", "last_commessa_hash"),
                    ("oggetto", "Oggetto", "rec_oggetto", "last_oggetto_hash")
                ]

                for campo_id, label, key_rec, key_hash in campi_tecnici:
                    with st.container():
                        # Registratore
                        audio_data = mic_recorder(key=key_rec, start_prompt=f"🎤 {label}", stop_prompt="🛑 FERMA E ANALIZZA")
                        
                        # Logica di elaborazione
                        if audio_data and isinstance(audio_data, dict) and 'bytes' in audio_data:
                            current_hash = hash(str(audio_data['bytes']))
                            if st.session_state.get(key_hash) != current_hash:
                                with st.spinner(f"Elaborazione {label}..."):
                                    set_bg_color("#D0AD00")
                                    
                                    risultato = elabora_anagrafica_ai(audio_data['bytes'], campo_id)
                                    st.session_state.anagrafica[campo_id] = risultato
                                    st.session_state[key_hash] = current_hash
                                    
                                    # INCREMENTIAMO IL SALT per forzare il refresh di TUTTI i widget
                                    st.session_state.anagrafica_version += 1
                                    
                                    salva_stato_completo()
                                    set_bg_color("#b3ff99")
                                    time.sleep(1)
                                    st.rerun()

                        # Widget con chiave dinamica basata sul salt
                        key_widget = f"widget_{campo_id}_{salt}"
                        
                        st.session_state.anagrafica[campo_id] = st.text_area(
                            label, 
                            value=st.session_state.anagrafica.get(campo_id, ""),
                            key=key_widget,
                            on_change=salva_stato_completo
                        )


        
        # EXPANDER 3: ATTIVITÀ E PERSONALE
        with st.expander("🛠️ Attività e Personale", expanded=True):
    
            with st.container():

                # Recuperiamo il salt corrente (se non esiste, partiamo da 0)
                salt = st.session_state.get("anagrafica_version", 0)

                # Lista di tuple per rendere il codice più compatto e DRY (Don't Repeat Yourself)
                campi_tecnici = [
                    ("attività", "Attività di Cantiere", "rec_attivita", "last_attivita_hash"),
                    ("coordinamento", "Coordinamento", "rec_coord", "last_coord_hash"),
                    ("personale", "Personale Presente", "rec_personale", "last_personale_hash"),
                    ("verbali", "Verbali di Prescrizione/Sospensione", "rec_verbali", "last_verb_hash")
                ]

                for campo_id, label, key_rec, key_hash in campi_tecnici:
                    with st.container():
                        # Registratore
                        audio_data = mic_recorder(key=key_rec, start_prompt=f"🎤 {label}", stop_prompt="🛑 FERMA E ANALIZZA")
                        
                        # Logica di elaborazione
                        if audio_data and isinstance(audio_data, dict) and 'bytes' in audio_data:
                            current_hash = hash(str(audio_data['bytes']))
                            if st.session_state.get(key_hash) != current_hash:
                                with st.spinner(f"Elaborazione {label}..."):
                                    set_bg_color("#D0AD00")
                                    
                                    risultato = elabora_campo_tecnico_ai(audio_data['bytes'], campo_id)
                                    st.session_state.anagrafica[campo_id] = risultato
                                    st.session_state[key_hash] = current_hash
                                    
                                    # INCREMENTIAMO IL SALT per forzare il refresh di TUTTI i widget
                                    st.session_state.anagrafica_version += 1
                                    
                                    salva_stato_completo()
                                    set_bg_color("#b3ff99")
                                    time.sleep(1)
                                    st.rerun()

                        # Widget con chiave dinamica basata sul salt
                        key_widget = f"widget_{campo_id}_{salt}"
                        
                        st.session_state.anagrafica[campo_id] = st.text_area(
                            label, 
                            value=st.session_state.anagrafica.get(campo_id, ""),
                            key=key_widget,
                            on_change=salva_stato_completo
                        )


        with st.expander("📎 Allegati"):
            uploaded_files = st.file_uploader(
                "Carica allegati", 
                accept_multiple_files=True, 
                type=['pdf', 'jpg', 'png', 'txt'],
                key="file_uploader_allegati",
                on_change=salva_stato_completo
            )
            
            if uploaded_files:
                nomi_file = ", ".join([f.name for f in uploaded_files])
                st.session_state.anagrafica["allegati"] = f"Elenco allegati: {nomi_file}"
                # Il salvataggio è già coperto dall'on_change del file_uploader
            else:
                # Se non ci sono file, salviamo lo stato di "vuoto"
                if st.session_state.anagrafica.get("allegati") != "Nessun allegato presente.":
                    st.session_state.anagrafica["allegati"] = "Nessun allegato presente."
                    salva_stato_completo()

            # Opzionale: mostra un avviso
            if st.session_state.anagrafica.get("allegati") != "Nessun allegato presente.":
                st.info(f"💾 {st.session_state.anagrafica['allegati']}")



        # --- ESPORTAZIONE ---
        st.divider()
        st.subheader("📥 Esportazione Report")
        
        if st.button("📄 Genera Report Finale"):
            with st.spinner("Generazione documento..."):
                doc_bytes = genera_report_finale(st.session_state.storico_report)
                st.session_state.doc_bytes = doc_bytes # Salviamo in session_state per l'email
                st.success("Report generato!")

        # Mostra i pulsanti solo se il report è stato generato
        if "doc_bytes" in st.session_state:
            col_down, col_mail = st.columns(2)
            
            with col_down:
                st.download_button(
                    label="✅ Scarica il report",
                    data=st.session_state.doc_bytes,
                    file_name="Report_Sicurezza.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            with col_mail:
                if st.button("📧 Invia a me stesso via Email"):
                    destinatario = st.session_state.user_data["email"]
                    with st.spinner(f"Invio email a {destinatario}..."):
                        success, msg = invia_report_via_email_graph(
                            st.session_state.doc_bytes, 
                            "Report_Sicurezza.docx", 
                            destinatario
                        )
                        if success:
                            st.success(msg)
                        else:
                            st.error(f"Errore: {msg}")